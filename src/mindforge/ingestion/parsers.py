"""Adaptive parsers for supported knowledge-base document formats."""

from __future__ import annotations

import hashlib
import html
import logging
import multiprocessing
import os
import re
import time
import threading
from collections.abc import Callable
from concurrent.futures import (
    ProcessPoolExecutor,
    ThreadPoolExecutor,
    as_completed,
)
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

logger = logging.getLogger(__name__)

ParserProgressCallback = Callable[[str, int, int], None]
ParserCancellationCallback = Callable[[], bool]
ElementKind = Literal["text", "table", "image"]


def _extract_pdf_page_range(
    path: str,
    page_indices: tuple[int, ...],
) -> list[tuple[int, str, str | None]]:
    """Extract a page range using a worker-local PDF handle."""
    import pdfplumber
    import warnings

    warnings.filterwarnings("ignore", message=".*FontBBox.*")
    warnings.filterwarnings("ignore", message=".*font descriptor.*")
    extracted: list[tuple[int, str, str | None]] = []
    with pdfplumber.open(path) as worker_pdf:
        for page_index in page_indices:
            try:
                text = worker_pdf.pages[page_index].extract_text() or ""
                extracted.append((page_index, text, None))
            except Exception as exc:
                extracted.append((page_index, "", f"{type(exc).__name__}: {exc}"))
    return extracted


class DocumentParserError(ValueError):
    """Expected document validation or parsing failure."""

    status_code = 422


class UnsupportedDocumentError(DocumentParserError):
    status_code = 400


class DocumentLimitError(DocumentParserError):
    status_code = 413


class DocumentParserCancelledError(DocumentParserError):
    """Cooperative cancellation at a parser page or model boundary."""

    status_code = 409


def _read_text_with_fallback(path: Path) -> str:
    """Read a text file with common Chinese and western encodings."""
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _sha256_path(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for block in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _to_plain_mapping(value: object) -> dict:
    """Return a Paddle result as a plain mapping across supported result forms."""
    if isinstance(value, dict):
        return value
    try:
        return dict(value)  # type: ignore[arg-type]
    except (TypeError, ValueError):
        pass
    json_method = getattr(value, "json", None)
    if callable(json_method):
        try:
            import json

            result = json_method
            payload = result() if callable(result) else result
            return json.loads(payload) if isinstance(payload, str) else dict(payload)
        except (TypeError, ValueError, AttributeError):
            pass
    return {}


def _as_list(value: object) -> list[object]:
    if value is None:
        return []
    if hasattr(value, "tolist"):
        value = value.tolist()  # type: ignore[union-attr]
    if isinstance(value, (list, tuple)):
        return list(value)
    return []


def _as_bbox(value: object) -> tuple[float, float, float, float] | None:
    points = _as_list(value)
    if not points:
        return None
    if len(points) == 4 and all(isinstance(point, (int, float)) for point in points):
        return tuple(float(point) for point in points)  # type: ignore[return-value]

    coordinates: list[tuple[float, float]] = []
    for point in points:
        pair = _as_list(point)
        if (
            len(pair) >= 4
            and all(isinstance(item, (int, float)) for item in pair[:4])
        ):
            coordinates.extend(
                [
                    (float(pair[0]), float(pair[1])),
                    (float(pair[2]), float(pair[3])),
                ]
            )
        elif (
            len(pair) >= 2
            and all(isinstance(item, (int, float)) for item in pair[:2])
        ):
            coordinates.append((float(pair[0]), float(pair[1])))
    if not coordinates:
        return None
    xs, ys = zip(*coordinates)
    return (min(xs), min(ys), max(xs), max(ys))


def _normalise_table_cell(value: object) -> str:
    text = "" if value is None else str(value)
    return re.sub(r"\s+", " ", text).strip().replace("|", r"\|")


def _table_to_structured(
    table: list[object],
    *,
    max_cells: int,
) -> tuple[str, str, list[dict], int, int] | None:
    rows: list[list[str]] = []
    merged_cells: list[list[bool]] = []
    cell_count = 0
    for raw_row in table:
        raw_cells = _as_list(raw_row)
        row = [_normalise_table_cell(cell) for cell in raw_cells]
        if not any(row):
            continue
        cell_count += len(row)
        if cell_count > max_cells:
            logger.warning(
                "Skipping table with %d cells; configured limit is %d.",
                cell_count,
                max_cells,
            )
            return None
        rows.append(row)
        merged_cells.append([cell is None for cell in raw_cells])
    if not rows:
        return None

    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    normalized_merged = [
        row + [False] * (column_count - len(row))
        for row in merged_cells
    ]
    header = normalized[0]
    separator = ["---"] * column_count
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
    cells = [
        {
            "row": row_index,
            "column": column_index,
            "text": value,
            "rowspan": 1,
            "colspan": 1,
            "header": row_index == 0,
            "is_merged": normalized_merged[row_index][column_index],
        }
        for row_index, row in enumerate(normalized)
        for column_index, value in enumerate(row)
    ]
    html_rows = []
    for row_index, row in enumerate(normalized):
        tag = "th" if row_index == 0 else "td"
        html_rows.append(
            "<tr>"
            + "".join(
                (
                    f"<{tag} data-merged=\""
                    f"{'true' if normalized_merged[row_index][column_index] else 'false'}"
                    f"\">{html.escape(value)}</{tag}>"
                )
                for column_index, value in enumerate(row)
            )
            + "</tr>"
        )
    return (
        "\n".join(lines),
        "<table>" + "".join(html_rows) + "</table>",
        cells,
        len(normalized),
        column_count,
    )


def _html_table_to_structured(
    html: str,
    *,
    max_cells: int,
) -> tuple[str, str, list[dict], int, int] | None:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    table = soup.find("table")
    if table is None:
        return None
    rows: list[list[str]] = []
    cells: list[dict] = []
    for row_index, row in enumerate(table.find_all("tr")):
        values: list[str] = []
        for column_index, cell in enumerate(
            row.find_all(["th", "td"], recursive=False)
        ):
            text = _normalise_table_cell(cell.get_text(" ", strip=True))
            values.append(text)
            cells.append(
                {
                    "row": row_index,
                    "column": column_index,
                    "text": text,
                    "rowspan": _positive_int(cell.get("rowspan"), 1),
                    "colspan": _positive_int(cell.get("colspan"), 1),
                    "header": cell.name == "th",
                }
            )
        if values:
            rows.append(values)
    converted = _table_to_structured(rows, max_cells=max_cells)
    if converted is None:
        return None
    markdown, _, _, row_count, column_count = converted
    return markdown, str(table), cells, row_count, column_count


def _positive_int(value: object, default: int) -> int:
    try:
        parsed = int(str(value))
    except (TypeError, ValueError):
        return default
    return parsed if parsed > 0 else default


def _looks_like_native_table(text: str) -> bool:
    lines = [line for line in text.splitlines() if line.strip()]
    if len(lines) < 3:
        return False
    aligned_rows = sum(
        1
        for line in lines
        if len(re.findall(r"\S(?:\s{2,}|\t)\S", line)) >= 1
    )
    return aligned_rows >= 2


def _looks_like_ocr_table(elements: list["DocumentElement"]) -> bool:
    text_elements = [element for element in elements if element.bbox is not None]
    if len(text_elements) < 6:
        return False
    left_edges = sorted(element.bbox[0] for element in text_elements if element.bbox)
    if not left_edges:
        return False
    clusters = [left_edges[0]]
    for edge in left_edges[1:]:
        if edge - clusters[-1] > 28:
            clusters.append(edge)
    return len(clusters) >= 3


def _average_confidence(elements: list["DocumentElement"]) -> float | None:
    values = [
        float(element.confidence)
        for element in elements
        if element.confidence is not None
    ]
    return sum(values) / len(values) if values else None


def _point_in_any_bbox(
    x: float,
    y: float,
    bboxes: list[tuple[float, float, float, float]],
) -> bool:
    return any(
        x0 <= x <= x1 and top <= y <= bottom
        for x0, top, x1, bottom in bboxes
    )


def _group_native_lines(words: list[dict[str, float | str]]) -> list[dict]:
    """Group word coordinates into visual lines without sharing PDF objects."""
    ordered = sorted(
        words,
        key=lambda word: (float(word["top"]), float(word["x0"])),
    )
    lines: list[list[dict[str, float | str]]] = []
    for word in ordered:
        if not lines:
            lines.append([word])
            continue
        current = lines[-1]
        current_top = sum(float(item["top"]) for item in current) / len(current)
        height = max(
            float(item["bottom"]) - float(item["top"]) for item in current
        )
        if abs(float(word["top"]) - current_top) <= max(2.5, height * 0.45):
            current.append(word)
        else:
            lines.append([word])
    output: list[dict] = []
    for line in lines:
        sorted_line = sorted(line, key=lambda word: float(word["x0"]))
        segments: list[list[dict[str, float | str]]] = [[sorted_line[0]]]
        for word in sorted_line[1:]:
            previous = segments[-1][-1]
            gap = float(word["x0"]) - float(previous["x1"])
            previous_width = max(
                float(previous["x1"]) - float(previous["x0"]),
                1.0,
            )
            if gap > max(32.0, previous_width * 3.0):
                segments.append([word])
            else:
                segments[-1].append(word)
        for segment in segments:
            output.append(
                {
                    "content": " ".join(
                        str(word["text"]) for word in segment
                    ).strip(),
                    "bbox": (
                        min(float(word["x0"]) for word in segment),
                        min(float(word["top"]) for word in segment),
                        max(float(word["x1"]) for word in segment),
                        max(float(word["bottom"]) for word in segment),
                    ),
                }
            )
    return output


def _group_native_blocks(lines: list[dict]) -> list[dict]:
    """Merge nearby lines into retrieval-sized blocks while preserving order."""
    if not lines:
        return []
    ordered = sorted(
        lines,
        key=lambda line: (line["bbox"][1], line["bbox"][0]),
    )
    columns = _order_lines_by_column(ordered)
    blocks: list[list[dict]] = []
    for column_lines in columns:
        if not column_lines:
            continue
        current_blocks: list[list[dict]] = [[column_lines[0]]]
        for line in column_lines[1:]:
            previous = current_blocks[-1][-1]["bbox"]
            bbox = line["bbox"]
            vertical_gap = bbox[1] - previous[3]
            horizontal_shift = abs(bbox[0] - previous[0])
            previous_width = max(previous[2] - previous[0], 1.0)
            if (
                vertical_gap <= 18.0
                and horizontal_shift <= max(72.0, previous_width * 0.45)
            ):
                current_blocks[-1].append(line)
            else:
                current_blocks.append([line])
        blocks.extend(current_blocks)
    output: list[dict] = []
    for block in blocks:
        bboxes = [line["bbox"] for line in block]
        output.append(
            {
                "content": "\n".join(
                    str(line["content"]) for line in block
                ).strip(),
                "bbox": (
                    min(bbox[0] for bbox in bboxes),
                    min(bbox[1] for bbox in bboxes),
                    max(bbox[2] for bbox in bboxes),
                    max(bbox[3] for bbox in bboxes),
                ),
                "line_count": len(block),
            }
        )
    return output


def _order_lines_by_column(lines: list[dict]) -> list[list[dict]]:
    """Return full-width material followed by each detected text column."""
    if len(lines) < 4:
        return [lines]
    starts = sorted(line["bbox"][0] for line in lines)
    gaps = [
        (right - left, (left + right) / 2)
        for left, right in zip(starts, starts[1:])
    ]
    if not gaps:
        return [lines]
    largest_gap, split = max(gaps, key=lambda item: item[0])
    if largest_gap < 72.0:
        return [lines]

    left = [
        line
        for line in lines
        if line["bbox"][0] < split and line["bbox"][2] <= split
    ]
    right = [
        line
        for line in lines
        if line["bbox"][0] >= split
    ]
    full_width = [
        line
        for line in lines
        if line not in left and line not in right
    ]
    if len(left) < 2 or len(right) < 2:
        return [lines]
    return [
        sorted(full_width, key=lambda line: (line["bbox"][1], line["bbox"][0])),
        sorted(left, key=lambda line: line["bbox"][1]),
        sorted(right, key=lambda line: line["bbox"][1]),
    ]


@dataclass
class DocumentElement:
    """A typed, source-aware unit emitted by the document parser."""

    kind: ElementKind
    content: str
    page: int | None = None
    bbox: tuple[float, float, float, float] | None = None
    confidence: float | None = None
    source_method: str = "native_text"
    metadata: dict = field(default_factory=dict)
    start: int | None = None
    end: int | None = None


@dataclass
class ParsedDocument:
    doc_id: str
    filename: str
    content: str
    sections: list[dict] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    images: list[dict] = field(default_factory=list)
    elements: list[DocumentElement] = field(default_factory=list)


class _PaddleOCRAdapter:
    """Lazy, process-local PaddleOCR adapter with result normalization."""

    _engines: dict[tuple[str, str, str, bool], object] = {}
    _lock = threading.Lock()

    def __init__(
        self,
        *,
        language: str,
        device: str,
        model_source: str,
        enable_mkldnn: bool,
    ) -> None:
        self._key = (language, device, model_source, enable_mkldnn)

    def _get_engine(self) -> object:
        with self._lock:
            cached = self._engines.get(self._key)
            if cached is not None:
                return cached
            language, device, model_source, enable_mkldnn = self._key
            os.environ["PADDLE_PDX_MODEL_SOURCE"] = model_source
            try:
                from paddleocr import PaddleOCR
            except ImportError as exc:
                raise DocumentParserError(
                    "OCR runtime is unavailable. Check the PaddleOCR package "
                    "and its system library dependencies."
                ) from exc
            try:
                engine = PaddleOCR(
                    lang=language,
                    device=device,
                    enable_mkldnn=enable_mkldnn,
                    use_doc_orientation_classify=False,
                    use_doc_unwarping=False,
                    use_textline_orientation=False,
                )
            except Exception as exc:
                raise DocumentParserError(
                    "OCR model initialization failed. Check the OCR device, "
                    "model cache, and model download source."
                ) from exc
            self._engines[self._key] = engine
            return engine

    def extract(self, image: object, *, page: int) -> list[DocumentElement]:
        try:
            import numpy as np

            payload = np.asarray(image.convert("RGB"))  # type: ignore[union-attr]
            result = list(self._get_engine().predict(payload))  # type: ignore[attr-defined]
        except DocumentParserError:
            raise
        except Exception as exc:
            raise DocumentParserError(
                f"OCR processing failed on page {page}."
            ) from exc
        if not result:
            return []

        data = _to_plain_mapping(result[0])
        texts = _as_list(data.get("rec_texts"))
        scores = _as_list(data.get("rec_scores"))
        polygons = _as_list(data.get("dt_polys") or data.get("rec_polys"))
        elements: list[DocumentElement] = []
        for index, raw_text in enumerate(texts):
            text = str(raw_text).strip()
            if not text:
                continue
            raw_score = scores[index] if index < len(scores) else None
            confidence = (
                float(raw_score)
                if isinstance(raw_score, (int, float))
                else None
            )
            polygon = polygons[index] if index < len(polygons) else None
            elements.append(
                DocumentElement(
                    kind="text",
                    content=text,
                    page=page,
                    bbox=_as_bbox(polygon),
                    confidence=confidence,
                    source_method="ocr",
                    metadata={"ocr_line_index": index},
                )
            )
        elements.sort(
            key=lambda element: (
                element.bbox[1] if element.bbox is not None else float("inf"),
                element.bbox[0] if element.bbox is not None else float("inf"),
            )
        )
        return elements


class _PaddleTableAdapter:
    """Optional table recognition adapter loaded only for likely OCR tables."""

    _pipelines: dict[tuple[str, bool], object] = {}
    _lock = threading.Lock()

    def __init__(
        self,
        *,
        device: str,
        model_source: str,
        enable_mkldnn: bool,
    ) -> None:
        self._device = device
        self._model_source = model_source
        self._enable_mkldnn = enable_mkldnn

    def _get_pipeline(self) -> object:
        with self._lock:
            cache_key = (self._device, self._enable_mkldnn)
            cached = self._pipelines.get(cache_key)
            if cached is not None:
                return cached
            os.environ["PADDLE_PDX_MODEL_SOURCE"] = self._model_source
            try:
                from paddleocr import TableRecognitionPipelineV2
            except ImportError as exc:
                raise DocumentParserError(
                    "Table recognition runtime is unavailable. Check the "
                    "PaddleOCR package and its system library dependencies."
                ) from exc
            try:
                pipeline = TableRecognitionPipelineV2(
                    device=self._device,
                    enable_mkldnn=self._enable_mkldnn,
                )
            except Exception as exc:
                raise DocumentParserError(
                    "Table recognition model initialization failed."
                ) from exc
            self._pipelines[cache_key] = pipeline
            return pipeline

    def extract(
        self,
        image: object,
        *,
        page: int,
        max_cells: int,
    ) -> list[DocumentElement]:
        try:
            import numpy as np

            payload = np.asarray(image.convert("RGB"))  # type: ignore[union-attr]
            result = list(self._get_pipeline().predict(payload))  # type: ignore[attr-defined]
        except DocumentParserError:
            raise
        except Exception as exc:
            raise DocumentParserError(
                f"Table recognition failed on page {page}."
            ) from exc
        if not result:
            return []

        data = _to_plain_mapping(result[0])
        tables = _as_list(data.get("table_res_list"))
        elements: list[DocumentElement] = []
        for index, raw_table in enumerate(tables):
            table = _to_plain_mapping(raw_table)
            html = str(table.get("pred_html") or "").strip()
            converted = _html_table_to_structured(html, max_cells=max_cells)
            if converted is None:
                continue
            markdown, table_html, cells, rows, columns = converted
            cell_boxes = _as_list(table.get("cell_box_list"))
            elements.append(
                DocumentElement(
                    kind="table",
                    content=markdown,
                    page=page,
                    bbox=_as_bbox(cell_boxes),
                    source_method="ocr_table",
                    metadata={
                        "table_index": index,
                        "row_count": rows,
                        "column_count": columns,
                        "table_html": table_html,
                        "table_cells": cells,
                    },
                )
            )
        return elements


class DocumentParser:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".md", ".txt"}

    def __init__(self) -> None:
        from mindforge.config import get_settings

        settings = get_settings()
        self._limits = settings.api
        self._parser_config = settings.parser
        self._progress_callback: ParserProgressCallback | None = None
        self._cancellation_callback: ParserCancellationCallback | None = None

    def set_progress_callback(
        self,
        callback: ParserProgressCallback | None,
    ) -> None:
        """Attach a synchronous callback used by background indexing jobs."""
        self._progress_callback = callback

    def set_cancellation_callback(
        self,
        callback: ParserCancellationCallback | None,
    ) -> None:
        """Attach a synchronous, thread-safe cooperative cancellation check."""
        self._cancellation_callback = callback

    def parse(
        self,
        file_path: str | Path,
        *,
        progress_callback: ParserProgressCallback | None = None,
        cancellation_callback: ParserCancellationCallback | None = None,
    ) -> ParsedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedDocumentError(
                f"Unsupported document format: {suffix}. Supported formats: "
                f"{', '.join(sorted(self.SUPPORTED_EXTENSIONS))}."
            )
        st = path.stat()
        if (
            suffix in {".html", ".htm", ".md", ".txt"}
            and st.st_size > self._limits.max_text_file_mb * 1024 * 1024
        ):
            raise DocumentLimitError(
                "Text document exceeds the configured parser size limit."
            )

        previous_callback = self._progress_callback
        previous_cancellation_callback = self._cancellation_callback
        if progress_callback is not None:
            self._progress_callback = progress_callback
        if cancellation_callback is not None:
            self._cancellation_callback = cancellation_callback
        try:
            self._raise_if_cancelled()
            if suffix == ".pdf":
                content, sections, metadata, elements, images = (
                    self._parse_pdf_structured(path)
                )
            else:
                parser = self._get_parser(suffix)
                content, sections, metadata = parser(path)
                elements = self._elements_from_content(content)
                images = []
        finally:
            if progress_callback is not None:
                self._progress_callback = previous_callback
            if cancellation_callback is not None:
                self._cancellation_callback = previous_cancellation_callback

        if len(content) > self._limits.max_parsed_chars:
            raise DocumentLimitError(
                "Parsed document exceeds the configured character limit."
            )
        visual_only = False
        if not content.strip():
            from mindforge.config import get_settings

            visual_only = bool(
                get_settings().visual_retrieval.enabled
                and any(element.kind == "image" for element in elements)
            )
            if not visual_only:
                raise DocumentParserError(
                    "Document contains no indexable text after native and OCR "
                    "extraction. Enable and configure visual retrieval before "
                    "indexing an image-only document."
                )
        doc_id = (
            _sha256_path(path)[:24]
            if visual_only
            else hashlib.sha256(content.encode("utf-8")).hexdigest()[:24]
        )
        metadata.update(
            {
                "source": path.name,
                "file_type": suffix,
                "size_bytes": st.st_size,
                "element_count": len(elements),
                "visual_only": visual_only,
            }
        )
        logger.info("Parsed %s (%d characters).", path.name, len(content))
        return ParsedDocument(
            doc_id=doc_id,
            filename=path.name,
            content=content,
            sections=sections,
            metadata=metadata,
            images=images,
            elements=elements,
        )

    def _get_parser(self, suffix: str):
        parsers = {
            ".docx": self._parse_docx,
            ".html": self._parse_html,
            ".htm": self._parse_html,
            ".md": self._parse_markdown,
            ".txt": self._parse_text,
        }
        return parsers[suffix]

    def _emit_progress(self, stage: str, completed: int, total: int) -> None:
        callback = self._progress_callback
        if callback is not None:
            callback(stage, completed, max(total, 1))

    def _raise_if_cancelled(self) -> None:
        callback = self._cancellation_callback
        if callback is not None and callback():
            raise DocumentParserCancelledError("Document parsing was cancelled.")

    def _extract_pdf_text_pages(
        self,
        path: Path,
    ) -> tuple[list[tuple[int, str, str | None]], int]:
        import pdfplumber
        import warnings

        warnings.filterwarnings("ignore", message=".*FontBBox.*")
        warnings.filterwarnings("ignore", message=".*font descriptor.*")
        with pdfplumber.open(str(path)) as pdf:
            total = len(pdf.pages)
            if total > self._limits.max_pdf_pages:
                raise DocumentLimitError(
                    f"PDF contains {total} pages, exceeding the configured "
                    f"limit of {self._limits.max_pdf_pages} pages."
                )
            self._emit_progress("detecting", 0, total)
            if total <= self._limits.pdf_parallel_page_threshold:
                results: list[tuple[int, str, str | None]] = []
                for index, page in enumerate(pdf.pages):
                    self._raise_if_cancelled()
                    try:
                        results.append((index, page.extract_text() or "", None))
                    except Exception as exc:
                        results.append(
                            (index, "", f"{type(exc).__name__}: {exc}")
                        )
                    self._emit_progress("detecting", index + 1, total)
                return results, total

        workers = min(
            self._limits.pdf_parse_workers,
            os.cpu_count() or 4,
            total,
        )
        executor_name = self._limits.pdf_parse_executor
        logger.info(
            "PDF native-text pass: %d pages, %d %s workers.",
            total,
            workers,
            executor_name,
        )
        pages_per_worker = (total + workers - 1) // workers
        pages_per_worker = min(
            pages_per_worker,
            getattr(self._limits, "pdf_parse_batch_pages", 8),
        )
        page_ranges = [
            tuple(range(start, min(total, start + pages_per_worker)))
            for start in range(0, total, pages_per_worker)
        ]

        def run_parallel(executor_type: type[ProcessPoolExecutor] | type[ThreadPoolExecutor]):
            executor_kwargs: dict[str, object] = {"max_workers": workers}
            if executor_type is ProcessPoolExecutor:
                executor_kwargs["mp_context"] = multiprocessing.get_context("spawn")
            range_results: list[list[tuple[int, str, str | None]]] = []
            with executor_type(**executor_kwargs) as pool:  # type: ignore[arg-type]
                futures = [
                    pool.submit(
                        _extract_pdf_page_range,
                        str(path),
                        page_range,
                    )
                    for page_range in page_ranges
                ]
                completed_pages = 0
                for future in as_completed(futures):
                    self._raise_if_cancelled()
                    result = future.result()
                    range_results.append(result)
                    completed_pages += len(result)
                    self._emit_progress(
                        "detecting",
                        min(completed_pages, total),
                        total,
                    )
            return range_results

        try:
            executor_type = (
                ProcessPoolExecutor
                if executor_name == "process"
                else ThreadPoolExecutor
            )
            range_results = run_parallel(executor_type)
        except Exception:
            if executor_name != "process":
                raise
            logger.exception("PDF process parsing failed; retrying with threads.")
            range_results = run_parallel(ThreadPoolExecutor)

        results = [
            item
            for range_result in range_results
            for item in range_result
        ]
        results.sort(key=lambda item: item[0])
        self._emit_progress("detecting", total, total)
        return results, total

    def _parse_pdf(self, path: Path):
        """Legacy tuple API retained for existing callers and tests."""
        results, total = self._extract_pdf_text_pages(path)
        content_parts: list[str] = []
        sections: list[dict] = []
        total_chars = 0
        for index, text, error in results:
            if error:
                logger.warning("PDF page %d extraction failed: %s", index + 1, error)
            total_chars += len(text)
            if total_chars > self._limits.max_parsed_chars:
                raise DocumentLimitError(
                    "PDF text exceeds the configured character limit."
                )
            content_parts.append(text)
            sections.append(
                {"title": f"Page {index + 1}", "content": text, "level": 0}
            )
        return "\n".join(content_parts), sections, {"pages": total}

    def _parse_pdf_structured(
        self,
        path: Path,
    ) -> tuple[str, list[dict], dict, list[DocumentElement], list[dict]]:
        import pdfplumber

        results, total = self._extract_pdf_text_pages(path)
        config = self._parser_config
        native_char_count = sum(len(text) for _, text, _ in results)
        if native_char_count > self._limits.max_parsed_chars:
            raise DocumentLimitError(
                "PDF text exceeds the configured character limit."
            )
        native_text = {index: text for index, text, _ in results}
        ocr_pages = [
            index
            for index, text, _ in results
            if self._page_needs_ocr(text)
        ]
        if config.mode == "ocr":
            ocr_pages = list(range(total))
        elif config.mode == "native" or not config.ocr_enabled:
            ocr_pages = []
        if len(ocr_pages) > config.ocr_max_pages:
            raise DocumentLimitError(
                f"Document requires OCR on {len(ocr_pages)} pages, exceeding "
                f"the configured OCR limit of {config.ocr_max_pages} pages."
            )

        page_elements: dict[int, list[DocumentElement]] = {
            index: (
                []
                if config.layout_enabled
                else self._native_page_elements(
                    object(),
                    page=index + 1,
                    text=text,
                    excluded_bboxes=[],
                )
            )
            for index, text in native_text.items()
        }
        table_candidates = {
            index
            for index, text in native_text.items()
            if config.table_extraction_enabled and _looks_like_native_table(text)
        }
        visual_pages = set(ocr_pages) | table_candidates
        if config.image_extraction_enabled or config.layout_enabled:
            visual_pages.update(range(total))

        ocr_failures: list[str] = []
        ocr_adapter: _PaddleOCRAdapter | None = None
        table_adapter: _PaddleTableAdapter | None = None
        ocr_completed = 0
        table_completed = 0
        page_metrics: list[dict] = []
        self._emit_progress("ocr", 0, len(ocr_pages))
        self._emit_progress("table", 0, max(len(table_candidates), 1))

        with pdfplumber.open(str(path)) as pdf:
            for index in sorted(visual_pages):
                self._raise_if_cancelled()
                page_started = time.perf_counter()
                page = pdf.pages[index]
                page_number = index + 1
                table_elements: list[DocumentElement] = []
                if index in table_candidates:
                    self._raise_if_cancelled()
                    table_elements = self._native_table_elements(
                        page,
                        page=page_number,
                    )
                    table_completed += 1
                    self._emit_progress(
                        "table",
                        table_completed,
                        max(len(table_candidates), 1),
                    )

                native_elements = (
                    self._native_page_elements(
                        page,
                        page=page_number,
                        text=native_text[index],
                        excluded_bboxes=[
                            table.bbox
                            for table in table_elements
                            if table.bbox is not None
                        ],
                    )
                    if config.layout_enabled
                    else []
                )
                image_elements: list[DocumentElement] = []
                if config.image_extraction_enabled:
                    image_elements = self._page_image_elements(
                        page,
                        page=page_number,
                    )
                page_elements[index].extend(native_elements)
                page_elements[index].extend(table_elements)
                page_elements[index].extend(image_elements)

                route_reasons = self._page_route_reasons(
                    text=native_text[index],
                    ocr_required=index in ocr_pages,
                    has_images=bool(image_elements),
                )

                if index in ocr_pages:
                    try:
                        self._raise_if_cancelled()
                        rendered = page.to_image(
                            resolution=config.ocr_dpi
                        ).original
                        if ocr_adapter is None:
                            ocr_adapter = _PaddleOCRAdapter(
                                language=config.ocr_language,
                                device=config.ocr_device,
                                model_source=config.ocr_model_source,
                                enable_mkldnn=config.ocr_enable_mkldnn,
                            )
                        ocr_elements = ocr_adapter.extract(
                            rendered,
                            page=page_number,
                        )
                        if ocr_elements:
                            average_confidence = _average_confidence(ocr_elements)
                            if (
                                average_confidence is not None
                                and average_confidence
                                < getattr(
                                    config,
                                    "ocr_handwriting_confidence",
                                    0.62,
                                )
                            ):
                                route_reasons.append("low_ocr_confidence")
                            page_elements[index] = [
                                *ocr_elements,
                                *[
                                    element
                                    for element in page_elements[index]
                                    if element.kind != "text"
                                ],
                            ]
                            if (
                                config.layout_enabled
                                and config.table_extraction_enabled
                                and _looks_like_ocr_table(ocr_elements)
                            ):
                                self._emit_progress("layout", 0, 1)
                                self._raise_if_cancelled()
                                if table_adapter is None:
                                    table_adapter = _PaddleTableAdapter(
                                        device=config.ocr_device,
                                        model_source=config.ocr_model_source,
                                        enable_mkldnn=config.ocr_enable_mkldnn,
                                    )
                                try:
                                    ocr_tables = table_adapter.extract(
                                        rendered,
                                        page=page_number,
                                        max_cells=config.table_max_cells,
                                    )
                                except DocumentParserError:
                                    logger.warning(
                                        "OCR table recognition skipped on page %d.",
                                        page_number,
                                        exc_info=True,
                                    )
                                else:
                                    if ocr_tables:
                                        page_elements[index].extend(ocr_tables)
                                self._emit_progress("layout", 1, 1)
                        elif native_text[index].strip():
                            logger.warning(
                                "OCR returned no text for page %d; retaining native text.",
                                page_number,
                            )
                        if getattr(config, "asset_render_ocr_pages", False):
                            self._ensure_visual_page_asset(
                                page,
                                page_elements[index],
                                page=page_number,
                                route_reasons=route_reasons,
                            )
                    except DocumentParserError as exc:
                        if native_text[index].strip():
                            logger.warning(
                                "OCR failed on page %d; retaining native text: %s",
                                page_number,
                                exc,
                            )
                        else:
                            ocr_failures.append(str(exc))
                    finally:
                        ocr_completed += 1
                        self._emit_progress(
                            "ocr",
                            ocr_completed,
                            len(ocr_pages),
                        )

                for element in page_elements[index]:
                    element.metadata["routing_reasons"] = sorted(
                        set(
                            [
                                *element.metadata.get(
                                    "routing_reasons",
                                    [],
                                ),
                                *route_reasons,
                            ]
                        )
                    )
                    element.metadata["requires_visual_review"] = bool(
                        {
                            "formula_candidate",
                            "low_ocr_confidence",
                        }
                        & set(route_reasons)
                    )
                page_metrics.append(
                    {
                        "page": page_number,
                        "native_characters": len(native_text[index]),
                        "ocr_used": index in ocr_pages,
                        "table_count": sum(
                            1
                            for element in page_elements[index]
                            if element.kind == "table"
                        ),
                        "image_count": sum(
                            1
                            for element in page_elements[index]
                            if element.kind == "image"
                        ),
                        "routing_reasons": sorted(set(route_reasons)),
                        "processing_ms": round(
                            (time.perf_counter() - page_started) * 1000,
                            3,
                        ),
                    }
                )

        content, sections, elements = self._assemble_pdf_output(page_elements, total)
        if not content.strip() and ocr_pages and ocr_failures:
            raise DocumentParserError(
                "OCR could not extract text from image-only PDF pages. "
                + ocr_failures[0]
            )
        return (
            content,
            sections,
            {
                "pages": total,
                "ocr_pages": len(ocr_pages),
                "native_text_pages": total - len(ocr_pages),
                "table_count": sum(
                    1 for element in elements if element.kind == "table"
                ),
                "image_count": sum(
                    1 for element in elements if element.kind == "image"
                ),
                "page_metrics": page_metrics,
            },
            elements,
            [
                dict(element.metadata)
                for element in elements
                if element.kind == "image"
            ],
        )

    def _page_needs_ocr(self, text: str) -> bool:
        normalized = "".join(char for char in text if not char.isspace())
        if len(normalized) < self._parser_config.ocr_min_native_text_chars:
            return True
        printable = sum(char.isprintable() for char in normalized)
        ratio = printable / len(normalized) if normalized else 0.0
        return ratio < self._parser_config.ocr_min_printable_ratio

    @staticmethod
    def _native_page_elements(
        pdf_page: object,
        *,
        page: int,
        text: str,
        excluded_bboxes: list[tuple[float, float, float, float]],
    ) -> list[DocumentElement]:
        cleaned = text.strip()
        if not cleaned:
            return []
        extract_words = getattr(pdf_page, "extract_words", None)
        if not callable(extract_words):
            return [
                DocumentElement(
                    kind="text",
                    content=cleaned,
                    page=page,
                    source_method="native_text",
                )
            ]
        try:
            raw_words = extract_words(
                use_text_flow=True,
                keep_blank_chars=False,
            )
        except Exception:
            logger.warning(
                "Native layout extraction skipped on page %d.",
                page,
                exc_info=True,
            )
            raw_words = []
        words = [
            {
                "text": str(word.get("text") or "").strip(),
                "x0": float(word.get("x0") or 0.0),
                "top": float(word.get("top") or 0.0),
                "x1": float(word.get("x1") or 0.0),
                "bottom": float(word.get("bottom") or 0.0),
            }
            for word in raw_words
            if isinstance(word, dict) and str(word.get("text") or "").strip()
        ]
        words = [
            word
            for word in words
            if not _point_in_any_bbox(
                (word["x0"] + word["x1"]) / 2,
                (word["top"] + word["bottom"]) / 2,
                excluded_bboxes,
            )
        ]
        if not words:
            return []
        lines = _group_native_lines(words)
        blocks = _group_native_blocks(lines)
        return [
            DocumentElement(
                kind="text",
                content=block["content"],
                page=page,
                bbox=block["bbox"],
                source_method="native_text",
                metadata={
                    "layout_source": "pdf_words",
                    "reading_order": block_index,
                    "line_count": block["line_count"],
                },
            )
            for block_index, block in enumerate(blocks)
            if block["content"].strip()
        ]

    def _page_image_elements(
        self,
        pdf_page: object,
        *,
        page: int,
    ) -> list[DocumentElement]:
        raw_images = _as_list(getattr(pdf_page, "images", []))
        elements: list[DocumentElement] = []
        for index, raw_image in enumerate(
            raw_images[: self._parser_config.image_max_per_page]
        ):
            image = _to_plain_mapping(raw_image)
            bbox = _as_bbox(
                [
                    image.get("x0"),
                    image.get("top"),
                    image.get("x1"),
                    image.get("bottom"),
                ]
            )
            metadata = {
                "page": page,
                "image_index": index,
                "width": image.get("width"),
                "height": image.get("height"),
                "bits": image.get("bits"),
                "colorspace": str(image.get("colorspace") or ""),
            }
            elements.append(
                DocumentElement(
                    kind="image",
                    content="",
                    page=page,
                    bbox=bbox,
                    source_method="embedded_pdf_image",
                    metadata=metadata,
                )
            )
        if len(raw_images) > self._parser_config.image_max_per_page:
            logger.warning(
                "Page %d has %d images; retaining only the configured %d.",
                page,
                len(raw_images),
                self._parser_config.image_max_per_page,
            )
        return elements

    @staticmethod
    def _page_route_reasons(
        *,
        text: str,
        ocr_required: bool,
        has_images: bool,
    ) -> list[str]:
        reasons: list[str] = []
        if ocr_required:
            reasons.append("ocr_required")
        if has_images:
            reasons.append("embedded_visual")
        if re.search(
            r"[\u00b1\u00d7\u00f7\u2200-\u22ff\u27c0-\u27ef]",
            text,
        ):
            reasons.append("formula_candidate")
        return reasons

    @staticmethod
    def _ensure_visual_page_asset(
        pdf_page: object,
        elements: list[DocumentElement],
        *,
        page: int,
        route_reasons: list[str],
    ) -> None:
        if any(element.kind == "image" for element in elements):
            return
        width = float(getattr(pdf_page, "width", 0.0) or 0.0)
        height = float(getattr(pdf_page, "height", 0.0) or 0.0)
        if width <= 0 or height <= 0:
            return
        elements.append(
            DocumentElement(
                kind="image",
                content="",
                page=page,
                bbox=(0.0, 0.0, width, height),
                source_method="rendered_ocr_page",
                metadata={
                    "page": page,
                    "width": width,
                    "height": height,
                    "routing_reasons": sorted(
                        set(["rendered_ocr_page", *route_reasons])
                    ),
                },
            )
        )

    def _native_table_elements(
        self,
        pdf_page: object,
        *,
        page: int,
    ) -> list[DocumentElement]:
        try:
            find_tables = getattr(pdf_page, "find_tables", None)
            if callable(find_tables):
                candidates = list(find_tables())
                tables = [
                    (candidate.extract(), _as_bbox(getattr(candidate, "bbox", None)))
                    for candidate in candidates
                ]
            else:
                tables = [
                    (table, None)
                    for table in getattr(pdf_page, "extract_tables")()
                ]
        except Exception:
            logger.warning(
                "Native table extraction skipped on page %d.",
                page,
                exc_info=True,
            )
            return []

        elements: list[DocumentElement] = []
        for index, (table, bbox) in enumerate(tables):
            converted = _table_to_structured(
                _as_list(table),
                max_cells=self._parser_config.table_max_cells,
            )
            if converted is None:
                continue
            markdown, table_html, cells, rows, columns = converted
            elements.append(
                DocumentElement(
                    kind="table",
                    content=markdown,
                    page=page,
                    bbox=bbox,
                    source_method="native_table",
                    metadata={
                        "table_index": index,
                        "row_count": rows,
                        "column_count": columns,
                        "table_html": table_html,
                        "table_cells": cells,
                    },
                )
            )
        return elements

    @staticmethod
    def _assemble_pdf_output(
        page_elements: dict[int, list[DocumentElement]],
        total_pages: int,
    ) -> tuple[str, list[dict], list[DocumentElement]]:
        all_elements: list[DocumentElement] = []
        sections: list[dict] = []
        content_parts: list[str] = []
        offset = 0
        for index in range(total_pages):
            page_content_parts: list[str] = []
            ordered_elements = sorted(
                page_elements.get(index, []),
                key=lambda element: (
                    element.bbox[1]
                    if element.bbox is not None
                    else float("inf"),
                    element.bbox[0]
                    if element.bbox is not None
                    else float("inf"),
                    int(element.metadata.get("reading_order", 0)),
                    element.kind,
                ),
            )
            for element in ordered_elements:
                all_elements.append(element)
                if not element.content.strip():
                    continue
                if content_parts:
                    offset += 2
                element.start = offset
                element.end = offset + len(element.content)
                content_parts.append(element.content)
                page_content_parts.append(element.content)
                offset = element.end
            sections.append(
                {
                    "title": f"Page {index + 1}",
                    "content": "\n\n".join(page_content_parts),
                    "level": 0,
                }
            )
        return "\n\n".join(content_parts), sections, all_elements

    @staticmethod
    def _elements_from_content(content: str) -> list[DocumentElement]:
        if not content.strip():
            return []
        return [
            DocumentElement(
                kind="text",
                content=content,
                source_method="native_text",
                start=0,
                end=len(content),
            )
        ]

    def _parse_docx(self, path: Path):
        from zipfile import BadZipFile, ZipFile

        try:
            with ZipFile(path) as archive:
                infos = archive.infolist()
                if len(infos) > self._limits.max_docx_parts:
                    raise DocumentLimitError(
                        "DOCX package exceeds the configured part limit."
                    )
                expanded_size = sum(info.file_size for info in infos)
                if (
                    expanded_size
                    > self._limits.max_docx_uncompressed_mb * 1024 * 1024
                ):
                    raise DocumentLimitError(
                        "DOCX package exceeds the configured expanded-size limit."
                    )
        except BadZipFile as exc:
            raise DocumentParserError("Invalid DOCX package.") from exc

        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        content_parts: list[str] = []
        sections: list[dict] = []
        total_chars = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            total_chars += len(para.text)
            if total_chars > self._limits.max_parsed_chars:
                raise DocumentLimitError(
                    "DOCX text exceeds the configured character limit."
                )
            content_parts.append(para.text)
            if para.style.name.startswith("Heading"):
                try:
                    level_string = para.style.name[len("Heading") :].strip()
                    level = int(level_string.split()[0]) if level_string else 1
                except (ValueError, IndexError):
                    level = 1
                sections.append(
                    {"title": para.text, "content": para.text, "level": level}
                )
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    total_chars += len(row_text)
                    if total_chars > self._limits.max_parsed_chars:
                        raise DocumentLimitError(
                            "DOCX text exceeds the configured character limit."
                        )
                    content_parts.append(row_text)
        return "\n".join(content_parts), sections, {}

    def _parse_html(self, path: Path):
        from bs4 import BeautifulSoup

        raw = _read_text_with_fallback(path)
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return text, [], {"title": soup.title.string if soup.title else ""}

    def _parse_markdown(self, path: Path):
        content = _read_text_with_fallback(path)
        sections = []
        for line in content.split("\n"):
            match = re.match(r"^(#{1,6})\s", line)
            if match:
                sections.append(
                    {
                        "title": line.lstrip("#").strip(),
                        "content": "",
                        "level": len(match.group(1)),
                    }
                )
        return content, sections, {}

    @staticmethod
    def _parse_text(path: Path):
        return _read_text_with_fallback(path), [], {}


class DirectoryParser:
    def __init__(self, parser: DocumentParser | None = None):
        self.parser = parser or DocumentParser()

    def parse_directory(
        self,
        dir_path: str | Path,
        recursive: bool = True,
    ) -> list[ParsedDocument]:
        documents = []
        base = Path(dir_path)
        pattern = "**/*" if recursive else "*"
        for file_path in sorted(base.glob(pattern)):
            if file_path.suffix.lower() not in self.parser.SUPPORTED_EXTENSIONS:
                continue
            try:
                documents.append(self.parser.parse(file_path))
            except Exception as exc:
                logger.warning("Failed to parse %s: %s", file_path.name, exc)
        logger.info("Parsed %d documents from %s.", len(documents), base)
        return documents
